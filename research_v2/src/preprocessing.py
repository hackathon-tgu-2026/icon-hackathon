import docx
import pandas as pd
import os

def read_docx (file_path):
    doc = docx.Document(file_path)
    text = []
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text.append(cell_text)

    return '\n'.join(text).strip()
    
def load_all_resumes(cv_folder = 'data/CV'):
    resumes = {}
    for fname in os.listdir(cv_folder):
        if fname.endswith('.docx'):
            rid = int(fname.split('.')[0])
            resumes[rid] = read_docx(os.path.join(cv_folder,fname))
    return resumes

def load_vacancies(file_path = 'data/5_vacancies.csv'):
    df = pd.read_csv(file_path)
    df = df[['id', 'job_description', 'job_title']]
    return df